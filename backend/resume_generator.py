from typing import Dict, List, Optional
from datetime import datetime
import os

# PDF generation using fpdf (pure Python)
from fpdf import FPDF
"""
NOTE:
- DOCX generation uses `python-docx` (import name: `docx`).
- If `python-docx` is not installed in the current environment, we must not
  crash the entire chat flow. So we import it lazily inside generate_docx().
"""


class ResumeGenerator:
    """Generate polished HTML resume from collected data"""
    
    def generate_html(self, resume_data: Dict) -> str:
        """Generate professional HTML resume"""
        
        name = resume_data.get("name", "Your Name")
        email = resume_data.get("email", "")
        linkedin = resume_data.get("linkedin", "")
        portfolio = resume_data.get("portfolio", "")
        summary = resume_data.get("summary", "")
        education = resume_data.get("education", [])
        skills = resume_data.get("skills", {})
        experience = resume_data.get("experience", [])
        projects = resume_data.get("projects", [])
        achievements = resume_data.get("achievements", [])
        certifications = resume_data.get("certifications", [])
        
        contact_html = self._build_contact_section(email, linkedin, portfolio)
        summary_html = self._build_summary_section(summary)
        education_html = self._build_education_section(education)
        skills_html = self._build_skills_section(skills)
        experience_html = self._build_experience_section(experience)
        projects_html = self._build_projects_section(projects)
        achievements_html = self._build_achievements_section(achievements)
        certifications_html = self._build_certifications_section(certifications)
        missing_skills = self._suggest_missing_skills(skills)
        missing_skills_html = self._build_missing_skills_section(missing_skills) if missing_skills else ""
        
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume - {name}</title>
    <style>
        * {{margin:0; padding:0; box-sizing:border-box;}}
        body {{font-family:'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height:1.6; color:#333; background:#f5f5f5; padding:20px;}}
        .resume-container {{max-width:210mm; min-height:297mm; margin:0 auto; background:white; padding:40px; box-shadow:0 0 10px rgba(0,0,0,0.1);}}
        .header {{text-align:center; border-bottom:3px solid #2563eb; padding-bottom:20px; margin-bottom:30px;}}
        .header h1 {{font-size:32px; color:#1e40af; margin-bottom:10px; font-weight:700;}}
        .contact-info {{display:flex; justify-content:center; gap:20px; flex-wrap:wrap; font-size:14px; color:#666;}}
        .contact-info a {{color:#2563eb; text-decoration:none;}}
        .contact-info a:hover {{text-decoration:underline;}}
        .section {{margin-bottom:30px;}}
        .section-title {{font-size:20px; color:#1e40af; border-bottom:2px solid #2563eb; padding-bottom:5px; margin-bottom:15px; font-weight:600;}}
        .summary-text {{text-align:justify; line-height:1.8; color:#444;}}
        .education-item, .experience-item, .project-item {{margin-bottom:20px; padding-left:15px; border-left:3px solid #e5e7eb; padding-bottom:15px;}}
        .education-item:last-child, .experience-item:last-child, .project-item:last-child {{border-left:none; padding-left:0;}}
        .item-title {{font-weight:600; color:#1e40af; font-size:16px; margin-bottom:5px;}}
        .item-details {{color:#666; font-size:14px; margin-bottom:8px;}}
        .skills-container {{display:flex; gap:30px; flex-wrap:wrap;}}
        .skills-category {{flex:1; min-width:200px;}}
        .skills-category h4 {{font-size:16px; color:#1e40af; margin-bottom:10px;}}
        .skills-list {{display:flex; flex-wrap:wrap; gap:8px;}}
        .skill-tag {{background:#eff6ff; color:#1e40af; padding:5px 12px; border-radius:15px; font-size:13px; border:1px solid #bfdbfe;}}
        ul {{list-style:none; padding-left:0;}}
        ul li {{padding-left:20px; position:relative; margin-bottom:8px; color:#555;}}
        ul li:before {{content:"•"; color:#2563eb; font-weight:bold; position:absolute; left:0;}}
        .missing-skills {{background:#fef3c7; border-left:4px solid #f59e0b; padding:15px; margin-top:20px; border-radius:4px;}}
        .missing-skills h4 {{color:#92400e; margin-bottom:10px;}}
        .missing-skills p {{color:#78350f; font-size:14px;}}
        @media print {{body {{background:white; padding:0;}} .resume-container {{box-shadow:none; padding:20px;}}}}
    </style>
</head>
<body>
    <div class="resume-container">
        <div class="header">
            <h1>{self._escape_html(name)}</h1>
            {contact_html}
        </div>
        
        {summary_html}
        {education_html}
        {skills_html}
        {experience_html}
        {projects_html}
        {achievements_html}
        {certifications_html}
        {missing_skills_html}
    </div>
</body>
</html>
"""
        return html

    def _escape_html(self, text: str) -> str:
        if not text:
            return ""
        return (str(text)
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&#x27;"))

    def _build_contact_section(self, email: str, linkedin: str, portfolio: str) -> str:
        contacts = []
        if email:
            contacts.append(f'<a href="mailto:{email}">{self._escape_html(email)}</a>')
        if linkedin:
            contacts.append(f'<a href="{linkedin}" target="_blank">LinkedIn</a>')
        if portfolio:
            contacts.append(f'<a href="{portfolio}" target="_blank">Portfolio</a>')
        return f'<div class="contact-info">{" | ".join(contacts)}</div>' if contacts else ""

    def _build_summary_section(self, summary: str) -> str:
        if not summary:
            return ""
        return f"""
        <div class="section">
            <h2 class="section-title">Professional Summary</h2>
            <p class="summary-text">{self._escape_html(summary)}</p>
        </div>
        """

    def _build_education_section(self, education: List[Dict]) -> str:
        if not education:
            return ""
        items_html = []
        for edu in education:
            details = edu.get("details", "")
            if details:
                items_html.append(f'<div class="education-item"><div class="item-title">Education</div><div class="item-details">{self._escape_html(details)}</div></div>')
        if not items_html:
            return ""
        return f"""
        <div class="section">
            <h2 class="section-title">Education</h2>
            {"".join(items_html)}
        </div>
        """

    def _build_skills_section(self, skills: Dict) -> str:
        if not skills:
            return ""
        technical = skills.get("technical", [])
        soft = skills.get("soft", [])
        technical_html = ""
        soft_html = ""
        if technical:
            technical_tags = "".join([f'<span class="skill-tag">{self._escape_html(skill)}</span>' for skill in technical])
            technical_html = f'<div class="skills-category"><h4>Technical Skills</h4><div class="skills-list">{technical_tags}</div></div>'
        if soft:
            soft_tags = "".join([f'<span class="skill-tag">{self._escape_html(skill)}</span>' for skill in soft])
            soft_html = f'<div class="skills-category"><h4>Soft Skills</h4><div class="skills-list">{soft_tags}</div></div>'
        if not technical_html and not soft_html:
            return ""
        return f"""
        <div class="section">
            <h2 class="section-title">Skills</h2>
            <div class="skills-container">
                {technical_html}
                {soft_html}
            </div>
        </div>
        """

    def _build_experience_section(self, experience: List[Dict]) -> str:
        if not experience:
            return ""
        items_html = []
        for exp in experience:
            details = exp.get("details", "")
            if details:
                items_html.append(f'<div class="experience-item"><div class="item-title">Experience</div><div class="item-details">{self._escape_html(details)}</div></div>')
        if not items_html:
            return ""
        return f"""
        <div class="section">
            <h2 class="section-title">Experience</h2>
            {"".join(items_html)}
        </div>
        """

    def _build_projects_section(self, projects: List[Dict]) -> str:
        if not projects:
            return ""
        items_html = []
        for project in projects:
            details = project.get("details", "")
            if details:
                items_html.append(f'<div class="project-item"><div class="item-title">Project</div><div class="item-details">{self._escape_html(details)}</div></div>')
        if not items_html:
            return ""
        return f"""
        <div class="section">
            <h2 class="section-title">Projects</h2>
            {"".join(items_html)}
        </div>
        """

    def _build_achievements_section(self, achievements: List[str]) -> str:
        if not achievements:
            return ""
        items_html = "".join([f"<li>{self._escape_html(achievement)}</li>" for achievement in achievements])
        return f"""
        <div class="section">
            <h2 class="section-title">Achievements & Awards</h2>
            <ul>{items_html}</ul>
        </div>
        """

    def _build_certifications_section(self, certifications: List[str]) -> str:
        if not certifications:
            return ""
        items_html = "".join([f"<li>{self._escape_html(cert)}</li>" for cert in certifications])
        return f"""
        <div class="section">
            <h2 class="section-title">Certifications</h2>
            <ul>{items_html}</ul>
        </div>
        """

    def _build_missing_skills_section(self, missing_skills: List[str]) -> str:
        skills_list = ", ".join(missing_skills)
        return f"""
        <div class="missing-skills">
            <h4>💡 Skills to Consider Adding</h4>
            <p>Based on your profile, you might want to consider highlighting or developing these skills: <strong>{self._escape_html(skills_list)}</strong></p>
        </div>
        """

    def _suggest_missing_skills(self, skills: Dict) -> List[str]:
        if not skills:
            return []
        technical = [s.lower() for s in skills.get("technical", [])]
        suggestions = []
        essential_skills = {
            "git": ["version control", "github"],
            "python": ["data structures", "algorithms"],
            "javascript": ["html", "css", "frontend development"],
            "communication": ["presentation skills", "public speaking"]
        }
        for key, value_list in essential_skills.items():
            if any(key in skill for skill in technical):
                for val in value_list:
                    if not any(val in skill for skill in technical):
                        suggestions.append(val)
        common_missing = ["Problem Solving", "Team Collaboration", "Time Management", "Adaptability"]
        if suggestions:
            suggestions.extend(common_missing[:2])
        return list(set(suggestions[:5]))

    def generate_pdf(self, resume_data: dict, output_file: str = "resume.pdf") -> str:
        """Generate PDF from resume data using FPDF.
        Creates a professional PDF resume with proper formatting.
        """
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)

            # Name and Contact Info
            name = resume_data.get("name", "Your Name")
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, name, ln=True, align="C")
            pdf.ln(5)

            # Contact information
            email = resume_data.get("email", "")
            linkedin = resume_data.get("linkedin", "")
            portfolio = resume_data.get("portfolio", "")

            contact_parts = []
            if email:
                contact_parts.append(f"Email: {email}")
            if linkedin:
                contact_parts.append(f"LinkedIn: {linkedin}")
            if portfolio:
                contact_parts.append(f"Portfolio: {portfolio}")

            if contact_parts:
                contact_text = " | ".join(contact_parts)
                pdf.set_font("Arial", "", 10)
                pdf.cell(0, 8, contact_text, ln=True, align="C")
                pdf.ln(5)

            # Summary
            summary = resume_data.get("summary", "")
            if summary:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Professional Summary", ln=True)
                pdf.set_font("Arial", "", 10)
                pdf.multi_cell(0, 6, summary)
                pdf.ln(3)

            # Skills
            skills = resume_data.get("skills", {})
            if skills:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Skills", ln=True)
                pdf.set_font("Arial", "", 10)
                for category, skill_list in skills.items():
                    if skill_list:
                        category_text = f"{category.title()}: {', '.join(skill_list)}"
                        pdf.cell(0, 6, category_text, ln=True)
                pdf.ln(3)

            # Experience
            experience = resume_data.get("experience", [])
            if experience:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Professional Experience", ln=True)
                pdf.set_font("Arial", "", 10)
                for exp in experience:
                    if isinstance(exp, dict):
                        title = exp.get("title", "")
                        company = exp.get("company", "")
                        duration = exp.get("duration", "")
                        description = exp.get("description", "")

                        if title or company:
                            header = f"{title}"
                            if company:
                                header += f" at {company}"
                            if duration:
                                header += f" ({duration})"
                            pdf.set_font("Arial", "B", 10)
                            pdf.cell(0, 6, header, ln=True)
                            pdf.set_font("Arial", "", 10)

                        if description:
                            # Handle bullet points
                            desc_lines = description.split('\n')
                            for line in desc_lines:
                                line = line.strip()
                                if line.startswith('•') or line.startswith('-'):
                                    pdf.cell(10, 5, "", ln=False)
                                    pdf.cell(0, 5, f"• {line[1:].strip()}", ln=True)
                                elif line:
                                    pdf.cell(0, 5, line, ln=True)
                        pdf.ln(2)
                pdf.ln(3)

            # Education
            education = resume_data.get("education", [])
            if education:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Education", ln=True)
                pdf.set_font("Arial", "", 10)
                for edu in education:
                    if isinstance(edu, dict):
                        degree = edu.get("degree", "")
                        school = edu.get("school", "")
                        year = edu.get("year", "")

                        edu_text = f"{degree}"
                        if school:
                            edu_text += f" - {school}"
                        if year:
                            edu_text += f" ({year})"

                        pdf.cell(0, 6, edu_text, ln=True)
                pdf.ln(3)

            # Projects
            projects = resume_data.get("projects", [])
            if projects:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Projects", ln=True)
                pdf.set_font("Arial", "", 10)
                for project in projects:
                    if isinstance(project, dict):
                        name = project.get("name", "")
                        description = project.get("description", "")
                        technologies = project.get("technologies", "")

                        if name:
                            pdf.set_font("Arial", "B", 10)
                            pdf.cell(0, 6, name, ln=True)
                            pdf.set_font("Arial", "", 10)

                        if technologies:
                            pdf.set_font("Arial", "I", 9)
                            pdf.cell(0, 5, f"Technologies: {technologies}", ln=True)
                            pdf.set_font("Arial", "", 10)

                        if description:
                            desc_lines = description.split('\n')
                            for line in desc_lines:
                                line = line.strip()
                                if line.startswith('•') or line.startswith('-'):
                                    pdf.cell(10, 5, "", ln=False)
                                    pdf.cell(0, 5, f"• {line[1:].strip()}", ln=True)
                                elif line:
                                    pdf.cell(0, 5, line, ln=True)
                        pdf.ln(2)
                pdf.ln(3)

            # Certifications
            certifications = resume_data.get("certifications", [])
            if certifications:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Certifications", ln=True)
                pdf.set_font("Arial", "", 10)
                for cert in certifications:
                    if cert:
                        pdf.cell(10, 5, "", ln=False)
                        pdf.cell(0, 5, f"• {cert}", ln=True)
                pdf.ln(3)

            # Achievements
            achievements = resume_data.get("achievements", [])
            if achievements:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "Achievements", ln=True)
                pdf.set_font("Arial", "", 10)
                for achievement in achievements:
                    if achievement:
                        pdf.cell(10, 5, "", ln=False)
                        pdf.cell(0, 5, f"• {achievement}", ln=True)
                pdf.ln(3)

            # Save the PDF
            pdf.output(output_file)

        except Exception as e:
            print("⚠️ PDF generation failed:", e)
            # Fallback: create a simple text file
            try:
                with open(output_file.replace('.pdf', '.txt'), 'w') as f:
                    f.write(f"Resume for {resume_data.get('name', 'User')}\n")
                    f.write("Generated with FPDF\n")
            except Exception:
                pass

        return output_file

    def generate_docx(self, resume_data: dict, output_file: str = "resume.docx") -> str:
        """Minimal DOCX generation using python-docx."""
        try:
            from docx import Document  # pip install python-docx
        except Exception:
            # If python-docx isn't installed, skip DOCX generation gracefully.
            return output_file

        doc = Document()
        name = resume_data.get("name", "Your Name")
        email = resume_data.get("email", "")
        linkedin = resume_data.get("linkedin", "")
        portfolio = resume_data.get("portfolio", "")
        doc.add_heading(name, level=0)
        contact_parts = [p for p in [email, linkedin, portfolio] if p]
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))
        summary = resume_data.get("summary")
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)
        education = resume_data.get("education") or []
        if education:
            doc.add_heading("Education", level=1)
            for edu in education:
                details = edu.get("details")
                if details:
                    doc.add_paragraph(details, style="List Bullet")
        skills = resume_data.get("skills") or {}
        technical = skills.get("technical") or []
        soft = skills.get("soft") or []
        if technical or soft:
            doc.add_heading("Skills", level=1)
            if technical:
                doc.add_paragraph("Technical: " + ", ".join(technical))
            if soft:
                doc.add_paragraph("Soft: " + ", ".join(soft))
        experience = resume_data.get("experience") or []
        if experience:
            doc.add_heading("Experience", level=1)
            for exp in experience:
                details = exp.get("details")
                if details:
                    doc.add_paragraph(details, style="List Bullet")
        projects = resume_data.get("projects") or []
        if projects:
            doc.add_heading("Projects", level=1)
            for project in projects:
                details = project.get("details")
                if details:
                    doc.add_paragraph(details, style="List Bullet")
        achievements = resume_data.get("achievements") or []
        certifications = resume_data.get("certifications") or []
        if achievements or certifications:
            doc.add_heading("Achievements & Certifications", level=1)
            for item in achievements + certifications:
                doc.add_paragraph(str(item), style="List Bullet")
        doc.save(output_file)
        return output_file
